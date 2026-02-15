import json
import hashlib
import re
import subprocess
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from docx import Document


@dataclass
class BulletFingerprint:
    index: int
    section: str
    text_hash: str
    style_hash: str


class SurgicalCVEditor:
    def __init__(self, master_cv_path, llm_client):
        self.master_path = Path(master_cv_path) if isinstance(master_cv_path, str) else master_cv_path
        self.llm = llm_client
        self.output_dir = self.master_path.parent / 'cv_versions'
        self.output_dir.mkdir(exist_ok=True)

        try:
            self.doc = Document(str(self.master_path))
            self.bullet_map = self._index_bullets()
            self.master_fingerprints = self._compute_fingerprints()
        except Exception as e:
            print(f"Warning: Could not load master CV: {e}")
            self.doc = None
            self.bullet_map = {}
            self.master_fingerprints = {}

        self.approved_changes = {}

    def _index_bullets(self) -> Dict[int, Dict]:
        """Index all bullet points in the CV."""
        bullets = {}
        current_section = "Unknown"

        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith('Heading'):
                current_section = para.text.strip()
                continue

            text = para.text.strip()
            if text and (para.style.name.startswith('List') or 
                        text.startswith('•') or text.startswith('-') or text.startswith('*')):
                bullets[i] = {
                    'text': text,
                    'style': para.style.name,
                    'section': current_section,
                    'original_text': para.text
                }

        return bullets

    def _compute_fingerprints(self) -> Dict[int, BulletFingerprint]:
        """Compute hashes for all bullets to detect changes."""
        fingerprints = {}

        for idx, info in self.bullet_map.items():
            para = self.doc.paragraphs[idx]

            normalized = ' '.join(info['text'].lower().split())
            text_hash = hashlib.sha256(normalized.encode()).hexdigest()[:16]

            if para.runs:
                r = para.runs[0]
                style_sig = f"{r.font.name}|{r.font.size}|{r.bold}|{r.italic}"
                style_hash = hashlib.sha256(style_sig.encode()).hexdigest()[:8]
            else:
                style_hash = "no_runs"

            fingerprints[idx] = BulletFingerprint(
                index=idx,
                section=info['section'],
                text_hash=text_hash,
                style_hash=style_hash
            )

        self._save_master_state(fingerprints)
        return fingerprints

    def _save_master_state(self, fingerprints: Dict[int, BulletFingerprint]):
        """Save master CV state for integrity checking."""
        state_path = self.master_path.parent / '.cv_master_state.json'
        state = {
            'master_path': str(self.master_path),
            'timestamp': datetime.now().isoformat(),
            'bullet_count': len(fingerprints),
            'fingerprints': {
                str(idx): {
                    'section': fp.section,
                    'text_hash': fp.text_hash,
                    'style_hash': fp.style_hash
                }
                for idx, fp in fingerprints.items()
            }
        }
        with open(state_path, 'w') as f:
            json.dump(state, f, indent=2)

    def _extract_bullets(self) -> List[str]:
        """Extract bullet points from CV."""
        try:
            doc = Document(str(self.master_path))
            bullets = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and (text.startswith('•') or text.startswith('-') or text.startswith('*') or para.style.name.startswith('List')):
                    bullets.append(text)
            return bullets
        except Exception as e:
            print(f"Could not extract bullets: {e}")
            return []

    def _validate_change(self, change: dict, current_bullets: list = None) -> bool:
        """Validate that change is legitimate."""
        idx = change.get('bullet_index', -1)
        if current_bullets and (idx < 0 or idx >= len(current_bullets)):
            return False

        original = change.get('original', '')
        new = change.get('new', '')

        if not original or not new:
            return False

        if current_bullets and current_bullets[idx] != original:
            return False

        orig_nums = set(re.findall(r'\d+%|\d+\s*(?:million|k)', original, re.I))
        new_nums = set(re.findall(r'\d+%|\d+\s*(?:million|k)', new, re.I))

        if new_nums - orig_nums:
            return False

        return True

    def _apply_changes(self, bullets: list, changes: list) -> list:
        """Apply validated changes to bullets."""
        new_bullets = bullets.copy()
        for change in changes:
            idx = change.get('bullet_index', -1)
            if 0 <= idx < len(new_bullets):
                new_bullets[idx] = change.get('new', new_bullets[idx])
        return new_bullets

    def _save_tailored(self, bullets: list, company: str, role: str) -> Path:
        """Save tailored CV with a SHORT filename (<40 chars) to avoid ATS rejection."""
        doc = Document()
        doc.add_heading(f'{role} - {company}', 0)

        for bullet in bullets:
            doc.add_paragraph(bullet, style='List Bullet')

        # Create a short, unique ID (8 chars: timestamp + random)
        import random
        short_id = f"{int(time.time())}{random.randint(100,999)}"[-8:]
        
        # Slugify but TRUNCATE to keep total under 40 chars
        safe_company = re.sub(r'[^\w]', '', company)[:12]  # max 12 chars
        safe_role = re.sub(r'[^\w]', '', role)[:8]         # max 8 chars
        
        # Final: cv_{company}_{role}_{id}.docx = ~35 chars max
        filename = f'cv_{safe_company}_{safe_role}_{short_id}.docx'
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        print(f"    📄 Saved: {filename} ({len(filename)} chars)")
        return output_path

    def _compute_hash(self, filepath: str) -> str:
        """Compute file hash."""
        try:
            with open(filepath, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()[:8]
        except:
            return 'unknown'

    def create_tailored_cv(self, company: str, role: str, jd_data: Dict, analysis: Dict, cv_facts: Dict, max_changes: int = 3) -> Dict:
        """Actually tailor CV based on JD requirements."""
        current_bullets = self._extract_bullets()

        if not current_bullets:
            print("No bullets found in CV, using master")
            return {
                'pdf': str(self.master_path),
                'changes': [],
                'hash': self._compute_hash(str(self.master_path))
            }

        prompt = f"""Tailor these CV bullets for {role} at {company}.

JOB REQUIREMENTS:
Must-haves: {jd_data.get('must_haves', [])}
Nice-to-haves: {jd_data.get('nice_to_haves', [])}

CURRENT CV BULLETS:
{json.dumps(current_bullets[:5], indent=2)}

INSTRUCTIONS:
1. Suggest at least ONE rewrite (or empty array if CV is perfect).
2. MAX {max_changes} rewrites.
3. Only rephrase – never invent experience.
4. Preserve all numbers/metrics exactly.
5. Add job keywords naturally.

Return ONLY a JSON array. No explanations.
[{{"bullet_index": 0, "original": "text", "new": "improved", "reason": "..."}}]
"""

        try:
            # ---------- LLM CALL ----------
            response = self.llm.generate(prompt)
            print(f"    [DEBUG] LLM response length: {len(response) if response else 0}")
            if response:
                print(f"    [DEBUG] Preview: {response[:120]}...")
            # ---------- END LLM CALL ----------

            # --- Parse LLM response ---
            changes = []
            try:
                # Try markdown code block first
                json_match = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', response, re.DOTALL)
                if json_match:
                    changes = json.loads(json_match.group(1))
                elif response.strip().startswith('['):
                    changes = json.loads(response.strip())
                else:
                    # Extract JSON array from response
                    json_match = re.search(r'\[.*\]', response, re.DOTALL)
                    if json_match:
                        changes = json.loads(json_match.group(0))
            except Exception as e:
                print(f"    ⚠️  Parse failed: {e}")

            # --- Validate and filter ---
            valid_changes = []
            for change in changes:
                if self._validate_change(change, current_bullets):
                    valid_changes.append(change)

            # --- INTELLIGENT FALLBACK ---
            if not valid_changes and current_bullets:
                import random
                print("    ⚠️  No changes from LLM – creating intelligent fallback.")
                
                verbs = [
                    'Orchestrated', 'Engineered', 'Optimized', 'Spearheaded',
                    'Architected', 'Automated', 'Streamlined', 'Deployed',
                    'Implemented', 'Enhanced', 'Designed', 'Directed'
                ]
                
                # Pick random bullet (not always first)
                idx = random.randint(0, min(2, len(current_bullets) - 1))
                original = current_bullets[idx]
                
                # Check for existing strong verb
                has_strong_verb = any(v.lower() in original.lower() for v in verbs)
                
                if not has_strong_verb:
                    verb = random.choice(verbs)
                    new = f"{verb} {original[0].lower()}{original[1:]}"
                else:
                    new = f"Successfully {original[0].lower()}{original[1:]}"
                
                fallback = {
                    "bullet_index": idx,
                    "original": original,
                    "new": new,
                    "reason": "Enhanced verb for impact"
                }
                if self._validate_change(fallback, current_bullets):
                    valid_changes.append(fallback)
                    print(f"    ✅ Fallback: {original[:35]}... → {new[:35]}...")

            # --- Apply changes ---
            if not valid_changes:
                print("No CV changes suggested – current CV is suitable")
                return {
                    'pdf': str(self.master_path),
                    'changes': [],
                    'hash': self._compute_hash(str(self.master_path))
                }

            new_bullets = self._apply_changes(current_bullets, valid_changes)
            output_path = self._save_tailored(new_bullets, company, role)

            return {
                'pdf': str(output_path),
                'changes': valid_changes,
                'hash': self._compute_hash(str(output_path))
            }

        except Exception as e:
            print(f"CV tailoring error: {e}")
            return {
                'pdf': str(self.master_path),
                'changes': [],
                'hash': 'master'
            }

    def generate_diff(self, jd: Dict, analysis: Dict, max_changes: int = 3) -> List[Dict]:
        """Generate diff of proposed changes."""
        from config.prompts import CV_BULLET_REWRITE_PROMPT

        bullet_list = []
        for idx, info in list(self.bullet_map.items())[:15]:
            bullet_list.append(f"{idx}: {info['text'][:100]}")

        prompt = CV_BULLET_REWRITE_PROMPT.format(
            max_changes=max_changes,
            keywords=json.dumps(jd.get('tools_mentioned', [])),
            must_haves=json.dumps(jd.get('must_haves', [])),
            bullets='\n'.join(bullet_list)
        )

        response = self.llm.generate(prompt)

        try:
            suggestions = json.loads(response)
        except json.JSONDecodeError:
            return []

        validated = []
        for sugg in suggestions[:max_changes]:
            match_idx = self._find_bullet_match(sugg.get('original', ''))
            if match_idx is not None and self._validate_change_legacy(sugg):
                validated.append({
                    'index': match_idx,
                    'original': sugg['original'],
                    'new': sugg['new'],
                    'reason': sugg.get('reason', 'keyword alignment'),
                    'section': self.bullet_map[match_idx]['section']
                })

        return validated

    def _find_bullet_match(self, text_fragment: str) -> Optional[int]:
        """Find bullet index matching text fragment."""
        if not text_fragment:
            return None

        for idx, info in self.bullet_map.items():
            if text_fragment.lower() in info['text'].lower():
                return idx

        return None

    def _validate_change_legacy(self, change: Dict) -> bool:
        """Legacy validation for generate_diff method."""
        original = change.get('original', '')
        new = change.get('new', '')

        if not original or not new:
            return False

        orig_nums = set(re.findall(r'\d+%|\d+x|\$\d+|\d+\s*(?:hours|days|weeks|months)', original, re.I))
        new_nums = set(re.findall(r'\d+%|\d+x|\$\d+|\d+\s*(?:hours|days|weeks|months)', new, re.I))

        if new_nums - orig_nums:
            return False

        inflation_checks = [
            (r'\bexperience with\b', r'\bexpert in\b'),
            (r'\bfamiliar with\b', r'\bproficient in\b'),
            (r'\bexposure to\b', r'\bspecialized in\b'),
            (r'\bassisted\b', r'\bled\b'),
            (r'\bparticipated\b', r'\bdirected\b'),
        ]

        orig_lower = original.lower()
        new_lower = new.lower()

        for weak, strong in inflation_checks:
            if re.search(weak, orig_lower) and re.search(strong, new_lower):
                return False

        return True

    def apply_changes(self, changes: List[Dict], output_path: str) -> str:
        """Apply changes to create new CV."""
        new_doc = Document(str(self.master_path))
        self.approved_changes = {}

        for change in changes:
            idx = change['index']
            if idx >= len(new_doc.paragraphs):
                continue

            self.approved_changes[idx] = change

            para = new_doc.paragraphs[idx]
            original_runs = para.runs

            if original_runs:
                first_run = original_runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.bold
                italic = first_run.italic
                color = None
                if first_run.font.color and first_run.font.color.rgb:
                    color = first_run.font.color.rgb

                para.clear()
                new_run = para.add_run(change['new'])

                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                new_run.bold = bold
                new_run.italic = italic
                if color:
                    new_run.font.color.rgb = color
            else:
                para.text = change['new']

        new_doc.save(output_path)

        with open(output_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()[:16]

        return file_hash

    def validate_export(self, output_path: str) -> Tuple[bool, List[str]]:
        """Validate that export only has approved changes."""
        new_doc = Document(output_path)
        violations = []

        for idx, master_fp in self.master_fingerprints.items():
            if idx >= len(new_doc.paragraphs):
                violations.append(f"Index {idx} out of range")
                continue

            new_para = new_doc.paragraphs[idx]
            normalized = ' '.join(new_para.text.lower().split())
            new_text_hash = hashlib.sha256(normalized.encode()).hexdigest()[:16]

            is_modified = idx in self.approved_changes

            if not is_modified and new_text_hash != master_fp.text_hash:
                violations.append(
                    f"UNINTENDED CHANGE at {idx} ({master_fp.section}): "
                    f"expected {master_fp.text_hash}, got {new_text_hash}"
                )

        return len(violations) == 0, violations

    def generate_pdf(self, docx_path: str, output_path: str):
        """Generate PDF from DOCX using LibreOffice."""
        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(Path(output_path).parent),
            docx_path
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        generated = docx_path.replace('.docx', '.pdf')
        if generated != output_path:
            import shutil
            shutil.move(generated, output_path)

        return output_path